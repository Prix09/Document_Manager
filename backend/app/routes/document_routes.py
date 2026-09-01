from fastapi import APIRouter, UploadFile, File, HTTPException
import os
import docx
from app.services.pdf_loader import extract_text_from_pdf
from app.services.text_cleaner import clean_text
from app.services.chunker import chunk_text
from app.services.embeddings import embed_text
from app.services.vector_store import vector_store
from app.models.document_models import DocumentUploadResponse

router = APIRouter(tags=["Documents"])

UPLOAD_DIR = "uploads/"
os.makedirs(UPLOAD_DIR, exist_ok=True)

def extract_text(file_path: str, filename: str) -> str:
    ext = filename.lower().split('.')[-1]
    if ext == 'pdf':
        return extract_text_from_pdf(file_path)
    elif ext == 'txt':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == 'docx':
        try:
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            # Fallback or clear error if the docx is invalid/corrupted
            raise ValueError(f"Invalid or corrupted DOCX file: {e}")
    else:
        raise ValueError(f"Unsupported file extension: {ext}")

@router.get("/list")
async def list_documents():
    if not os.path.exists(UPLOAD_DIR):
        return {"documents": []}
    files = [f for f in os.listdir(UPLOAD_DIR) if os.path.isfile(os.path.join(UPLOAD_DIR, f))]
    return {"documents": files}

from typing import List

import re

def sanitize_filename(filename: str) -> str:
    # Keep alphanumeric, dot, underscore, dash, and space
    sanitized = re.sub(r'[^a-zA-Z0-9.\_\-\s]', '_', filename)
    # Trim to 100 characters to prevent path length issues
    if len(sanitized) > 100:
        name, ext = os.path.splitext(sanitized)
        sanitized = name[:100 - len(ext)] + ext
    return sanitized

from fastapi import APIRouter, UploadFile, File, HTTPException, BackgroundTasks

def process_file_in_background(file_path: str, filename: str, safe_filename: str):
    print(f"DEBUG (BG): Extracting text from {safe_filename}")
    try:
        text = extract_text(file_path, filename)
        print(f"DEBUG (BG): Extracted {len(text)} characters")
        cleaned = clean_text(text)
        chunks = chunk_text(cleaned)
        print(f"DEBUG (BG): Created {len(chunks)} chunks")
        
        if chunks:
            vectors = embed_text(chunks)
            vector_store.add_vectors(vectors, chunks, safe_filename)
            print(f"DEBUG (BG): Successfully embedded and stored {safe_filename}")
    except Exception as e:
        print(f"DEBUG (BG): Failed to process {safe_filename}: {e}")

@router.post("/", response_model=DocumentUploadResponse)
async def upload_document(background_tasks: BackgroundTasks, files: List[UploadFile] = File(...)):
    for file in files:
        safe_filename = sanitize_filename(file.filename)
        print(f"DEBUG: Receiving upload for {ascii(file.filename)} -> {safe_filename}")
        
        file_path = os.path.join(UPLOAD_DIR, safe_filename)

        # Save file synchronously to ensure it exists before background task starts
        try:
            with open(file_path, "wb") as buffer:
                content = await file.read()
                buffer.write(content)
            print(f"DEBUG: Successfully saved {ascii(file.filename)} to {file_path}")
        except Exception as e:
            print(f"DEBUG: Error saving file: {str(e)}")
            raise HTTPException(status_code=500, detail="Failed to save file")

        # Offload the heavy embedding task to the background
        background_tasks.add_task(process_file_in_background, file_path, file.filename, safe_filename)

    return DocumentUploadResponse(
        message=f"{len(files)} Document(s) uploading and processing in background!",
        chunks=0
    )
